from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
import fitz
from docx import Document
from pathlib import Path
import uuid

app = FastAPI(title="PDF to DOCX Service")

BASE_DIR = Path(__file__).parent
UPLOADS_DIR = BASE_DIR / "uploads"
OUTPUTS_DIR = BASE_DIR / "outputs"

UPLOADS_DIR.mkdir(exist_ok=True)
OUTPUTS_DIR.mkdir(exist_ok=True)

def convert_pdf_to_docx(pdf_path: str, docx_path: str):
    pdf = fitz.open(pdf_path)
    document = Document()

    for page in pdf:
        text = page.get_text("text")
        if text.strip():
            for line in text.split("\n"):
                document.add_paragraph(line)

    document.save(docx_path)


def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    doc = Document(docx_path)
    pdf = fitz.open()
    page = pdf.new_page()

    y = 50
    for para in doc.paragraphs:
        page.insert_text((50, y), para.text, fontsize=11)
        y += 15

        if y > 800:
            page = pdf.new_page()
            y = 50

    pdf.save(pdf_path)



@app.post("/convert/pdf-to-docx/")
async def pdf_to_docx_api(file: UploadFile = File(...)):
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files allowed")

    file_id = str(uuid.uuid4())
    pdf_path = UPLOADS_DIR / f"{file_id}.pdf"
    docx_path = OUTPUTS_DIR / f"{file_id}.docx"

    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    convert_pdf_to_docx(str(pdf_path), str(docx_path))

    return FileResponse(
        path=docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="converted.docx"
    )

@app.post("/convert/docx-to-pdf/")
async def docx_to_pdf_api(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only DOCX files allowed")

    file_id = str(uuid.uuid4())
    docx_path = UPLOADS_DIR / f"{file_id}.docx"
    pdf_path = OUTPUTS_DIR / f"{file_id}.pdf"

    with open(docx_path, "wb") as f:
        f.write(await file.read())

    convert_docx_to_pdf(str(docx_path), str(pdf_path))

    return FileResponse(
        path=pdf_path,
        media_type="application/pdf",
        filename="converted.pdf"
    )

