import fitz
from docx import Document
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
OUTPUTS_DIR = BASE_DIR / "outputs"

OUTPUTS_DIR.mkdir(exist_ok=True)

def convert_pdf_to_docx(pdf_path: str, docx_path: str):
    pdf = fitz.open(pdf_path)
    document = Document()

    for page in pdf:
        text = page.get_text("text")
        if text.strip():
            for line in text.split("\n"):
                document.add_paragraph(line)

    document.save(docx_path)


def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    doc = Document(docx_path)
    pdf = fitz.open()
    page = pdf.new_page()

    y = 50
    for para in doc.paragraphs:
        page.insert_text((50, y), para.text, fontsize=11)
        y += 15

        if y > 800:
            page = pdf.new_page()
            y = 50

    pdf.save(pdf_path)


    